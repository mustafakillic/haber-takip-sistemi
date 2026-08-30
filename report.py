"""Günlük Harekât Durum Raporu (GHDR) — Word (.docx) üreteci.

Analist, tarama sonuçlarından haberleri seçer ve her biri için 5N1K
(Kim / Ne / Nerede / Ne zaman / Neden / Nasıl) sorularını tek bir paragrafta
yanıtlayan özeti girer; rapora yalnızca seçilen haberler ve bu paragraflar
dökülür. Rapora bağlantı yazılmaz.

Kapsam kuralı: bir önceki gün 08:00'dan raporun oluşturulduğu ana kadar
yayımlanan haberler.
"""

import io
from datetime import datetime, timedelta

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

CLASSIFICATION = "HİZMETE ÖZEL"

_TR_MONTHS = [
    "OCA", "ŞUB", "MAR", "NİS", "MAY", "HAZ",
    "TEM", "AĞU", "EYL", "EKİ", "KAS", "ARA",
]


def report_period(now: datetime | None = None):
    """(başlangıç, bitiş) döner.
    Bitiş = raporun oluşturulduğu an ("Rapor Oluştur"a basılan zaman).
    Başlangıç = bir önceki takvim gününün saat 08:00'i.
    Örn. 30 Ağustos 14:20'de basılırsa: 29 Ağustos 08:00 → 30 Ağustos 14:20."""
    now = now or datetime.now()
    end = now.replace(second=0, microsecond=0)
    start = (end - timedelta(days=1)).replace(hour=8, minute=0, second=0, microsecond=0)
    return start, end


def _fmt(dt: datetime) -> str:
    return f"{dt.day:02d} {_TR_MONTHS[dt.month - 1]} {dt.year} – {dt.strftime('%H:%M')}"


def _report_no(end: datetime) -> str:
    return f"GHDR-{end.year}/{end.strftime('%m%d')}-01"


def filter_window(items, start: datetime, end: datetime):
    """search_raw() çıktısını kapsanan döneme göre süzer.
    Seçim ekranında en yeni haber üstte olacak şekilde yeniden eskiye sıralar
    (rapordaki kronolojik döküm ayrıca eskiden yeniye sıralanır)."""
    out = []
    for it in items:
        dt = it.get("sort_key")
        if not isinstance(dt, datetime) or dt == datetime.min:
            continue
        if start <= dt <= end:
            out.append(it)
    out.sort(key=lambda r: r["sort_key"], reverse=True)
    return out


def _add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = k
        cells[1].text = v
        for p in cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
    return table


def _news_block(doc, it, prefix=""):
    """Bir haberi: kalın başlık satırı + tek paragraflık 5N1K özeti olarak yazar."""
    head = doc.add_paragraph()
    head.add_run(f"{prefix}{it['title']} ").bold = True
    head.add_run(f"({it['source']} · {it['published']})").italic = True
    ozet = (it.get("ozet") or "").strip() or "[5N1K özeti girilmemiştir.]"
    doc.add_paragraph(ozet)


def build_report(selected, now: datetime | None = None, quick_keywords=None) -> bytes:
    """selected: analiste seçili haberler.
        [{ 'category', 'title', 'source', 'published', 'sort_key' (opsiyonel), 'ozet' }]
    'ozet' 5N1K sorularını tek paragrafta yanıtlayan analist metnidir.
    """
    quick_keywords = quick_keywords or []
    start, end = report_period(now)
    generated = now or datetime.now()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    cls = doc.add_paragraph(CLASSIFICATION)
    cls.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cls.runs[0].bold = True

    title = doc.add_heading("GÜNLÜK HAREKÂT DURUM RAPORU (GHDR)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Kars Bölgesi Harekât Merkezi — Sabah Arzı")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    note = doc.add_paragraph()
    note.add_run(
        "Kaynak niteliği: Bu rapor açık kaynak (basın / RSS / kurumsal duyuru) "
        "taramasına dayanır. Her haber, 5N1K (Kim / Ne / Nerede / Ne zaman / Neden / "
        "Nasıl) sorularını tek paragrafta yanıtlayacak şekilde analist tarafından "
        "özetlenmiştir; değerlendirmeler teyitli istihbarat hükmünde değildir."
    ).italic = True

    # kategori sırasını koru
    categories = []
    for it in selected:
        if it["category"] not in categories:
            categories.append(it["category"])

    sources = {it["source"] for it in selected if it.get("source")}

    # 1. RAPOR KİMLİĞİ
    doc.add_heading("1. RAPOR KİMLİĞİ", level=1)
    _add_kv_table(doc, [
        ("Rapor No", _report_no(end)),
        ("Kapsanan Dönem", f"{_fmt(start)} / {_fmt(end)}"),
        ("Kapsam Kuralı", "Bir önceki gün 08:00'dan raporun oluşturulduğu ana kadar açık kaynakta yayımlanan haberler"),
        ("Hazırlama Zamanı", _fmt(generated)),
        ("Hazırlayan", "Harekât Merkezi Amirliği"),
        ("Gizlilik Derecesi", CLASSIFICATION),
        ("Rapora Alınan Haber / Kaynak", f"{len(selected)} haber / {len(sources)} kaynak"),
    ])

    # 2. YÖNETİCİ ÖZETİ
    doc.add_heading("2. YÖNETİCİ ÖZETİ", level=1)
    flash = [
        it for it in selected
        if any(k.lower() in (it["title"] + " " + (it.get("ozet") or "")).lower() for k in quick_keywords)
    ]
    durum = "DİKKAT" if flash else ("OLAĞAN" if selected else "SAKİN")
    doc.add_paragraph(f"Rapora alınan haber sayısı: {len(selected)}", style="List Bullet")
    doc.add_paragraph(
        "Konu başlığı dağılımı: "
        + (", ".join(f"{c} ({sum(1 for it in selected if it['category'] == c)})" for c in categories) or "—"),
        style="List Bullet",
    )
    doc.add_paragraph(f"Kritik/öncelikli anahtar kelime içeren madde sayısı: {len(flash)}", style="List Bullet")
    doc.add_paragraph(f"Genel durum değerlendirmesi: {durum}", style="List Bullet")
    doc.add_paragraph("Serbest metin durum değerlendirmesi: ____________________________________")

    # 3. KRİTİK / ÖNCELİKLİ OLAYLAR
    doc.add_heading("3. KRİTİK / ÖNCELİKLİ OLAYLAR", level=1)
    if flash:
        for idx, it in enumerate(flash, 1):
            _news_block(doc, it, prefix=f"{idx}. [{it['category']}] ")
    else:
        doc.add_paragraph("KAYDA DEĞER GELİŞME YOKTUR.")

    # 4. KONU BAŞLIKLARINA GÖRE DURUM (KBİ)
    doc.add_heading("4. KONU BAŞLIKLARINA GÖRE DURUM (KBİ)", level=1)
    if not selected:
        doc.add_paragraph("Rapora haber seçilmemiştir.")
    for i, cat in enumerate(categories, 1):
        rows = [it for it in selected if it["category"] == cat]
        doc.add_heading(f"4.{i}. {cat} ({len(rows)})", level=2)
        for it in rows:
            _news_block(doc, it)

    # 5. OLAY DÖKÜMÜ (Kronolojik) — bağlantısız
    doc.add_heading("5. OLAY DÖKÜMÜ (KRONOLOJİK)", level=1)
    chron = sorted(selected, key=lambda it: it.get("sort_key") or datetime.min)
    if chron:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for i, h in enumerate(["Zaman", "Konu Başlığı", "Olay", "Kaynak"]):
            t.rows[0].cells[i].text = h
        for it in chron:
            c = t.add_row().cells
            c[0].text = it["published"]
            c[1].text = it["category"]
            c[2].text = it["title"].strip()
            c[3].text = it["source"]
    else:
        doc.add_paragraph("Kayıt yoktur.")

    # 6. KOMUTAN KARARINA / TALİMATINA SUNULAN HUSUSLAR
    doc.add_heading("6. KOMUTAN KARARINA / TALİMATINA SUNULAN HUSUSLAR", level=1)
    doc.add_paragraph(
        "Yalnızca komutanın kararı, talimatı veya kaynak tahsisi gereken hususlar. "
        "Her madde: Husus – Değerlendirme – Teklif – Beklenen Karar sırasıyla yazılır. "
        "Karar gerektiren husus yoksa aşağıdaki ilk satır bırakılır."
    ).runs[0].italic = True
    doc.add_paragraph(
        "1. Bu raporlama döneminde komutan kararına sunulacak yeni bir husus "
        "bulunmamaktadır."
    )
    for n in (2, 3):
        p = doc.add_paragraph()
        p.add_run(f"{n}. Husus: ").bold = True
        p.add_run("____________________________________")
        for lbl in ("Değerlendirme", "Teklif", "Beklenen Karar"):
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Pt(20)
            sp.add_run(f"{lbl}: ").bold = True
            sp.add_run("____________________________________")

    doc.add_paragraph()
    sign = doc.add_paragraph("Hazırlayan: ...............   Kontrol: ...............   Arz: ...............")
    sign.alignment = WD_ALIGN_PARAGRAPH.CENTER

    foot = doc.add_paragraph(CLASSIFICATION)
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.runs[0].bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def report_filename(now: datetime | None = None) -> str:
    _, end = report_period(now)
    return f"GHDR_{end.strftime('%Y%m%d')}.docx"
